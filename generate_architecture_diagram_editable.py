"""
Generate the MHA-DQN Architecture Diagram in multiple editable formats:
- Word Document (.docx) with embedded diagram
- SVG (Scalable Vector Graphics) - editable in Inkscape, Illustrator, etc.
- PDF - editable in some tools
- High-resolution PNG (as reference)
"""

import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from pathlib import Path
from datetime import datetime

# Try to import docx for Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Word document will not be generated.")

def draw_architecture_diagram(fig, ax):
    """Draw the complete architecture diagram on the given figure and axes."""
    
    # Actual feature groups used in the model with all features
    feature_groups = [
        ("Group 1", "Price & Returns", ["return", "volatility", "sma_20", "sma_50", "rsi_14", "log_return", "price_change"]),
        ("Group 2", "Macroeconomic", ["inflation", "GDP_growth", "unemployment", "fed_funds_rate", "t10y2y", "m2_money_supply"]),
        ("Group 3", "Commodities", ["wti_crude", "gold", "silver", "natural_gas", "copper"]),
        ("Group 4", "Market Indices", ["spy", "qqq", "vix", "xlk", "xlf", "xle"]),
        ("Group 5", "Forex", ["eur_usd", "gbp_usd", "usd_jpy", "aud_usd"]),
        ("Group 6", "Technical", ["macd", "macd_signal", "bb_upper", "bb_lower", "atr", "adx", "momentum", "roc", "williams_r"]),
        ("Group 7", "Earnings/Sentiment", ["earnings_reported_eps", "earnings_surprise_pct", "days_since_earnings", "approaching_earnings", "sentiment_score"]),
        ("Group 8", "Crypto", ["btc_usd", "eth_usd", "xrp_usd"]),
    ]
    
    # Colors for each group
    colors = {
        'input': '#2563eb',
        'head': '#00d4aa',
        'layer': '#7c3aed',
        'output': '#f59e0b',
        'arrow': '#64748b',
        'title': '#1e40af',
        'explanation': '#475569'
    }
    
    # Vertical positions with generous spacing to avoid overlaps
    y_positions = {
        'title': 17.0,
        'input_title': 15.2,
        'input_explanation': 14.6,
        'input': 13.0,
        'input_to_head_gap': 0.6,
        'head_title': 11.0,
        'head_explanation': 10.4,
        'head': 9.5,
        'head_to_layer_gap': 0.6,
        'layer_title': 7.8,
        'layer_explanation': 7.2,
        'layers': 6.0,
        'layer_to_pool_gap': 0.6,
        'pool_title': 4.2,
        'pool_explanation': 3.6,
        'pooling': 2.8,
        'pool_to_out_gap': 0.6,
        'out_title': 1.4,
        'out_explanation': 0.8,
        'output': 0.0
    }
    
    # Helper function to draw simple rectangular box
    def draw_simple_box(x, y, width, height, color, edge_color, linewidth=2, alpha=0.9, zorder=2):
        """Draw a simple rectangular box with frame."""
        rect = mpatches.Rectangle((x, y), width, height, 
                                 facecolor=color, edgecolor=edge_color, 
                                 linewidth=linewidth, zorder=zorder, alpha=alpha)
        ax.add_patch(rect)
        return rect
    
    # Main title
    ax.text(10, y_positions['title'], 
           'MHA-DQN Architecture: Complete Data Flow', 
           ha='center', fontsize=16, weight='bold', color=colors['title'], 
           fontfamily='monospace')
    
    # ========== INPUT FEATURE GROUPS SECTION ==========
    ax.text(10, y_positions['input_title'], 'INPUT FEATURE GROUPS (8 Groups)', 
           ha='center', fontsize=12, weight='bold', color=colors['input'], 
           fontfamily='monospace')
    
    ax.text(10, y_positions['input_explanation'], 
           'Each group contains related features that are processed together', 
           ha='center', fontsize=9, style='italic', color=colors['explanation'], 
           fontfamily='monospace')
    
    # Calculate positions for input feature groups
    box_width = 2.4
    spacing = 0.25
    row_width = 4 * box_width + 3 * spacing
    x_start = (20 - row_width) / 2
    box_height = 2.0
    y_start = y_positions['input']
    
    input_box_info = []
    
    for i, (group_num, group_name, features) in enumerate(feature_groups):
        row = i // 4
        col = i % 4
        x = x_start + col * (box_width + spacing)
        y = y_start - row * (box_height + 0.4)
        
        box_center_x = x + box_width/2
        box_bottom_y = y
        input_box_info.append((box_center_x, box_bottom_y, x, y))
        
        # Draw simple box
        draw_simple_box(x, y, box_width, box_height, colors['input'], '#4299e1', linewidth=2)
        
        # Group number
        ax.text(x + box_width/2, y + box_height - 0.12, group_num, 
               ha='center', va='top', fontsize=11, weight='bold', color='white', fontfamily='monospace')
        
        # Group name
        ax.text(x + box_width/2, y + box_height - 0.28, group_name, 
               ha='center', va='top', fontsize=9, weight='600', color='#e0e7ff', fontfamily='monospace')
        
        # All features as vertical list
        y_text_start = y + box_height - 0.48
        line_height = 0.20
        
        for idx, feature in enumerate(features):
            y_pos = y_text_start - idx * line_height
            ax.text(x + 0.15, y_pos, '•', 
                   ha='left', va='center', fontsize=7, color='#c7d2fe', fontfamily='monospace')
            ax.text(x + 0.3, y_pos, feature, 
                   ha='left', va='center', fontsize=7, color='#c7d2fe', fontfamily='monospace')
    
    # ========== ATTENTION HEADS SECTION ==========
    ax.text(10, y_positions['head_title'], '8 ATTENTION HEADS (H1-H8)', 
           ha='center', fontsize=12, weight='bold', color=colors['head'], 
           fontfamily='monospace')
    
    ax.text(10, y_positions['head_explanation'], 
           'Each head processes one feature group using Self-Attention (Q, K, V matrices)', 
           ha='center', fontsize=9, style='italic', color=colors['explanation'], 
           fontfamily='monospace')
    
    # Draw attention heads box
    head_box_width = 14
    head_box_height = 1.0
    head_y = y_positions['head']
    head_x = 3
    
    draw_simple_box(head_x, head_y, head_box_width, head_box_height, colors['head'], '#00d4aa', linewidth=2.5)
    
    ax.text(10, head_y + head_box_height/2 + 0.25, 'Self-Attention Mechanism: Q (Query), K (Key), V (Value)', 
           ha='center', va='center', fontsize=10, weight='bold', color='white', fontfamily='monospace')
    ax.text(10, head_y + head_box_height/2 - 0.05, 'Each head independently analyzes relationships within its feature group', 
           ha='center', va='center', fontsize=8, color='#d1fae5', fontfamily='monospace')
    ax.text(10, head_y + head_box_height/2 - 0.25, 'H1→Group1, H2→Group2, H3→Group3, H4→Group4, H5→Group5, H6→Group6, H7→Group7, H8→Group8', 
           ha='center', va='center', fontsize=7.5, color='#a7f3d0', fontfamily='monospace')
    
    # Arrows from input to heads
    center_x = 10
    for i in range(8):
        x_input_center, y_input_bottom, _, _ = input_box_info[i]
        y_head_top = head_y + head_box_height
        dy = y_head_top - y_input_bottom + y_positions['input_to_head_gap']
        dx = center_x - x_input_center
        ax.arrow(x_input_center, y_input_bottom, dx, dy,
                head_width=0.15, head_length=0.12, fc=colors['arrow'], ec=colors['arrow'],
                linewidth=1.5, zorder=1, alpha=0.7, length_includes_head=True)
    
    # ========== ATTENTION LAYERS SECTION ==========
    ax.text(10, y_positions['layer_title'], '3 ATTENTION LAYERS (Stacked)', 
           ha='center', fontsize=12, weight='bold', color=colors['layer'], 
           fontfamily='monospace')
    
    ax.text(10, y_positions['layer_explanation'], 
           'Each layer: Multi-Head Attention → Layer Norm → Feed-Forward (128→512→128) → Layer Norm', 
           ha='center', fontsize=9, style='italic', color=colors['explanation'], 
           fontfamily='monospace')
    
    # Draw 3 attention layers box
    layer_box_width = 14
    layer_box_height = 1.2
    layer_y = y_positions['layers']
    layer_x = 3
    
    draw_simple_box(layer_x, layer_y, layer_box_width, layer_box_height, colors['layer'], '#7c3aed', linewidth=2.5, alpha=0.85)
    ax.text(10, layer_y + layer_box_height/2 + 0.35, 'Layer 1: Learns basic temporal relationships & short-term patterns', 
           ha='center', va='center', fontsize=9.5, weight='bold', color='#c4b5fd', fontfamily='monospace')
    ax.text(10, layer_y + layer_box_height/2 + 0.1, 'Multi-Head Attention (8 heads) → Layer Norm → Feed-Forward (128→512→128) → Layer Norm', 
           ha='center', va='center', fontsize=8, color='#a78bfa', fontfamily='monospace')
    ax.text(10, layer_y + layer_box_height/2 - 0.15, 'Residual connections preserve original information while adding learned patterns', 
           ha='center', va='center', fontsize=7.5, color='#ddd6fe', fontfamily='monospace', style='italic')
    
    # Layer 2 and 3 descriptions
    layer2_y = layer_y - 0.3
    ax.text(10, layer2_y, 'Layer 2: Builds on Layer 1 to capture medium-term dependencies', 
           ha='center', va='center', fontsize=9, weight='600', color='#7c3aed', fontfamily='monospace')
    layer3_y = layer2_y - 0.25
    ax.text(10, layer3_y, 'Layer 3: Integrates all information for long-term patterns & complex relationships', 
           ha='center', va='center', fontsize=9, weight='600', color='#7c3aed', fontfamily='monospace')
    
    # Arrow from heads to layers
    center_x = 10
    y_head_bottom = head_y
    y_layer_top = layer_y + layer_box_height
    dy = y_layer_top - y_head_bottom + y_positions['head_to_layer_gap']
    ax.arrow(center_x, y_head_bottom, 0, dy,
            head_width=0.18, head_length=0.15, fc=colors['arrow'], ec=colors['arrow'],
            linewidth=2.5, zorder=1)
    
    # ========== GLOBAL AVERAGE POOLING SECTION ==========
    ax.text(10, y_positions['pool_title'], 'GLOBAL AVERAGE POOLING', 
           ha='center', fontsize=12, weight='bold', color=colors['layer'], 
           fontfamily='monospace')
    
    ax.text(10, y_positions['pool_explanation'], 
           'Averages all time steps in the sequence to create a single 128-dimensional vector', 
           ha='center', fontsize=9, style='italic', color=colors['explanation'], 
           fontfamily='monospace')
    
    # Global Average Pooling box
    pool_width = 10
    pool_y = y_positions['pooling']
    pool_height = 0.7
    pool_x = 5
    
    draw_simple_box(pool_x, pool_y, pool_width, pool_height, colors['layer'], '#7c3aed', linewidth=2.5, alpha=0.85)
    ax.text(10, pool_y + pool_height/2 + 0.15, 'Global Average Pooling', 
           ha='center', va='center', fontsize=10, weight='bold', color='#c4b5fd', fontfamily='monospace')
    ax.text(10, pool_y + pool_height/2 - 0.15, 'Sequence (20 time steps × 128 dim) → Single Vector (128 dim)',
           ha='center', va='center', fontsize=8.5, color='#a78bfa', fontfamily='monospace')
    
    # Arrow from layers to pooling
    center_x = 10
    y_layer_bottom = layer_y
    y_pooling_top = pool_y + pool_height
    dy = y_pooling_top - y_layer_bottom + y_positions['layer_to_pool_gap']
    ax.arrow(center_x, y_layer_bottom, 0, dy,
            head_width=0.18, head_length=0.15, fc=colors['arrow'], ec=colors['arrow'],
            linewidth=2.5, zorder=1)
    
    # ========== OUTPUT SECTION ==========
    ax.text(10, y_positions['out_title'], 'OUTPUT: Q-VALUES', 
           ha='center', fontsize=12, weight='bold', color=colors['output'], 
           fontfamily='monospace')
    
    ax.text(10, y_positions['out_explanation'], 
           'Q-values represent expected future rewards for each action (BUY, HOLD, SELL)', 
           ha='center', fontsize=9, style='italic', color=colors['explanation'], 
           fontfamily='monospace')
    
    # Output: Q-Values box
    output_width = 8
    output_height = 0.7
    output_y = y_positions['output']
    output_x = 6
    
    draw_simple_box(output_x, output_y, output_width, output_height, colors['output'], '#f59e0b', linewidth=2.5)
    ax.text(10, output_y + output_height/2 + 0.15, '3 Actions: BUY • HOLD • SELL',
           ha='center', va='center', fontsize=10, weight='bold', color='white', fontfamily='monospace')
    ax.text(7.5, output_y + output_height/2 - 0.1, 'BUY', ha='center', va='center', fontsize=8, 
           weight='bold', color='#fef3c7', fontfamily='monospace')
    ax.text(10, output_y + output_height/2 - 0.1, 'HOLD', ha='center', va='center', fontsize=8,
           weight='bold', color='#fef3c7', fontfamily='monospace')
    ax.text(12.5, output_y + output_height/2 - 0.1, 'SELL', ha='center', va='center', fontsize=8,
           weight='bold', color='#fef3c7', fontfamily='monospace')
    
    # Arrow from pooling to output
    center_x = 10
    y_pooling_bottom = pool_y
    y_output_top = output_y + output_height
    dy = y_output_top - y_pooling_bottom + y_positions['pool_to_out_gap']
    ax.arrow(center_x, y_pooling_bottom, 0, dy,
            head_width=0.18, head_length=0.15, fc=colors['arrow'], ec=colors['arrow'],
            linewidth=2.5, zorder=1)
    
    ax.set_xlim(0, 20)
    ax.set_ylim(-0.5, 18.0)

def generate_all_formats(output_dir='.'):
    """Generate the architecture diagram in multiple formats."""
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    print("=" * 60)
    print("Generating MHA-DQN Architecture Diagram in Multiple Formats")
    print("=" * 60)
    
    # ========== 1. Generate PNG (High Resolution) ==========
    print("\n[1/4] Generating PNG (High Resolution)...")
    fig, ax = plt.subplots(figsize=(20, 18), facecolor='#f5f7fa')
    ax.set_facecolor('#f5f7fa')
    ax.axis('off')
    
    draw_architecture_diagram(fig, ax)
    
    png_path = output_path / 'MHA_DQN_Architecture_Diagram.png'
    plt.savefig(png_path, format='png', dpi=300, facecolor='#f5f7fa', edgecolor='none', bbox_inches='tight')
    plt.close()
    print(f"  ✓ PNG saved: {png_path}")
    
    # ========== 2. Generate SVG (Editable Vector Format) ==========
    print("\n[2/4] Generating SVG (Editable Vector Format)...")
    fig, ax = plt.subplots(figsize=(20, 18), facecolor='#f5f7fa')
    ax.set_facecolor('#f5f7fa')
    ax.axis('off')
    
    draw_architecture_diagram(fig, ax)
    
    svg_path = output_path / 'MHA_DQN_Architecture_Diagram.svg'
    plt.savefig(svg_path, format='svg', facecolor='#f5f7fa', edgecolor='none', bbox_inches='tight')
    plt.close()
    print(f"  ✓ SVG saved: {svg_path}")
    print(f"    → Editable in: Inkscape (free), Adobe Illustrator, Figma, etc.")
    
    # ========== 3. Generate PDF ==========
    print("\n[3/4] Generating PDF...")
    fig, ax = plt.subplots(figsize=(20, 18), facecolor='#f5f7fa')
    ax.set_facecolor('#f5f7fa')
    ax.axis('off')
    
    draw_architecture_diagram(fig, ax)
    
    pdf_path = output_path / 'MHA_DQN_Architecture_Diagram.pdf'
    plt.savefig(pdf_path, format='pdf', facecolor='#f5f7fa', edgecolor='none', bbox_inches='tight')
    plt.close()
    print(f"  ✓ PDF saved: {pdf_path}")
    print(f"    → Editable in: Adobe Acrobat Pro, Inkscape, etc.")
    
    # ========== 4. Generate Word Document with Embedded Diagram ==========
    if DOCX_AVAILABLE:
        print("\n[4/4] Generating Word Document (.docx)...")
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Title
        title = doc.add_heading('MHA-DQN Architecture Diagram', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph('Complete Data Flow: 8 Feature Groups → 8 Attention Heads → 3 Attention Layers')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.italic = True
        
        # Date
        date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()  # Spacing
        
        # Add the PNG image
        if png_path.exists():
            doc.add_picture(str(png_path), width=Inches(7.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Add detailed explanation section
        doc.add_heading('Architecture Explanation', 1)
        
        doc.add_heading('Input Feature Groups (8 Groups)', 2)
        doc.add_paragraph(
            'The architecture starts with 8 feature groups, each containing related financial and market indicators:'
        )
        
        feature_groups_text = [
            ("Group 1: Price & Returns", "return, volatility, sma_20, sma_50, rsi_14, log_return, price_change"),
            ("Group 2: Macroeconomic", "inflation, GDP_growth, unemployment, fed_funds_rate, t10y2y, m2_money_supply"),
            ("Group 3: Commodities", "wti_crude, gold, silver, natural_gas, copper"),
            ("Group 4: Market Indices", "spy, qqq, vix, xlk, xlf, xle"),
            ("Group 5: Forex", "eur_usd, gbp_usd, usd_jpy, aud_usd"),
            ("Group 6: Technical", "macd, macd_signal, bb_upper, bb_lower, atr, adx, momentum, roc, williams_r"),
            ("Group 7: Earnings/Sentiment", "earnings_reported_eps, earnings_surprise_pct, days_since_earnings, approaching_earnings, sentiment_score"),
            ("Group 8: Crypto", "btc_usd, eth_usd, xrp_usd"),
        ]
        
        for group_name, features in feature_groups_text:
            p = doc.add_paragraph(group_name, style='List Bullet')
            p.add_run(f': {features}').font.size = Pt(10)
        
        doc.add_heading('8 Attention Heads (H1-H8)', 2)
        doc.add_paragraph(
            'Each attention head processes one feature group independently using Self-Attention mechanism. '
            'The Self-Attention uses Query (Q), Key (K), and Value (V) matrices to analyze relationships '
            'between different time steps within each feature group. This allows the model to focus on '
            'the most relevant information from the historical sequence.'
        )
        
        doc.add_heading('3 Attention Layers (Stacked)', 2)
        doc.add_paragraph(
            'The three attention layers are stacked to create hierarchical feature learning:'
        )
        doc.add_paragraph('Layer 1: Learns basic temporal relationships and short-term patterns', style='List Bullet')
        doc.add_paragraph('Layer 2: Builds on Layer 1 to capture medium-term dependencies', style='List Bullet')
        doc.add_paragraph('Layer 3: Integrates all information for long-term patterns and complex relationships', style='List Bullet')
        
        doc.add_paragraph()
        doc.add_paragraph(
            'Each layer contains:',
            style='List Bullet'
        )
        doc.add_paragraph(
            '1. Multi-Head Attention: Processes all 8 heads together',
            style='List Bullet 2'
        )
        doc.add_paragraph(
            '2. Layer Normalization: Stabilizes training by normalizing activations',
            style='List Bullet 2'
        )
        doc.add_paragraph(
            '3. Feed-Forward Network: Expands dimensions (128→512) for complex pattern learning, then compresses back (512→128)',
            style='List Bullet 2'
        )
        doc.add_paragraph(
            '4. Layer Normalization: Final normalization before passing to next layer',
            style='List Bullet 2'
        )
        doc.add_paragraph(
            'Residual connections (skip connections) preserve original information while allowing the network to learn additive improvements.',
            style='List Bullet 2'
        )
        
        doc.add_heading('Global Average Pooling', 2)
        doc.add_paragraph(
            'Converts the sequence of 20 time steps (each with 128 dimensions) into a single 128-dimensional vector. '
            'This summarizes all temporal information into a fixed-size representation suitable for final decision-making.'
        )
        
        doc.add_heading('Output: Q-Values', 2)
        doc.add_paragraph(
            'The final output consists of Q-values for three actions: BUY, HOLD, and SELL. '
            'Each Q-value represents the expected future reward for taking that action in the current market state. '
            'The action with the highest Q-value is selected as the trading recommendation.'
        )
        
        # Save Word document
        docx_path = output_path / 'MHA_DQN_Architecture_Diagram.docx'
        doc.save(str(docx_path))
        print(f"  ✓ Word Document saved: {docx_path}")
        print(f"    → Editable in: Microsoft Word, Google Docs, LibreOffice, etc.")
    else:
        print("\n[4/4] Skipping Word Document (python-docx not available)")
        print("      Install with: pip install python-docx")
    
    print("\n" + "=" * 60)
    print("✅ All diagrams generated successfully!")
    print("=" * 60)
    print(f"\nOutput directory: {output_path.absolute()}")
    print("\nGenerated files:")
    print(f"  • {png_path.name} - High-resolution PNG (300 DPI)")
    print(f"  • {svg_path.name} - Scalable Vector Graphics (editable)")
    print(f"  • {pdf_path.name} - PDF document")
    if DOCX_AVAILABLE:
        print(f"  • {docx_path.name} - Word document with diagram and explanations")

if __name__ == "__main__":
    # Generate all formats in current directory
    generate_all_formats(output_dir='.')
    
    print("\n💡 Editing Tips:")
    print("  • SVG: Best for vector editing (Inkscape, Illustrator, Figma)")
    print("  • Word: Best for adding text explanations and annotations")
    print("  • PDF: Good for presentations and printing")

