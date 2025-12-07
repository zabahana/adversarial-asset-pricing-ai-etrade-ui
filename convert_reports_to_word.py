#!/usr/bin/env python3
"""
Convert markdown reports to Word documents (.docx).
"""

import re
from pathlib import Path
from typing import List
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def setup_document_styles(doc: Document):
    """Configure document styles for professional formatting."""
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Configure heading styles
    heading1 = doc.styles['Heading 1']
    heading1_font = heading1.font
    heading1_font.name = 'Calibri'
    heading1_font.size = Pt(16)
    heading1_font.bold = True
    heading1_font.color.rgb = RGBColor(0, 0, 0)
    
    heading2 = doc.styles['Heading 2']
    heading2_font = heading2.font
    heading2_font.name = 'Calibri'
    heading2_font.size = Pt(14)
    heading2_font.bold = True
    heading2_font.color.rgb = RGBColor(44, 62, 80)
    
    heading3 = doc.styles['Heading 3']
    heading3_font = heading3.font
    heading3_font.name = 'Calibri'
    heading3_font.size = Pt(12)
    heading3_font.bold = True
    heading3_font.color.rgb = RGBColor(52, 73, 94)


def parse_markdown_line(line: str) -> dict:
    """Parse a markdown line and return its type and content."""
    line = line.rstrip()
    
    # Empty line
    if not line.strip():
        return {'type': 'empty', 'content': ''}
    
    # Heading level 1
    if line.startswith('# '):
        return {'type': 'h1', 'content': line[2:].strip()}
    
    # Heading level 2
    if line.startswith('## '):
        return {'type': 'h2', 'content': line[3:].strip()}
    
    # Heading level 3
    if line.startswith('### '):
        return {'type': 'h3', 'content': line[4:].strip()}
    
    # Heading level 4
    if line.startswith('#### '):
        return {'type': 'h4', 'content': line[5:].strip()}
    
    # Code block start/end
    if line.strip() == '```':
        return {'type': 'code_block', 'content': ''}
    
    # Code block with language
    if line.startswith('```'):
        return {'type': 'code_block_start', 'content': line[3:].strip()}
    
    # Bullet point
    if line.startswith('- ') or line.startswith('* '):
        return {'type': 'bullet', 'content': line[2:].strip()}
    
    # Numbered list
    match = re.match(r'^\d+\.\s+(.+)$', line)
    if match:
        return {'type': 'numbered', 'content': match.group(1)}
    
    # Bold text
    if line.startswith('**') and line.endswith('**'):
        return {'type': 'bold', 'content': line[2:-2]}
    
    # Table of contents marker
    if line.strip() == '---' or line.strip() == '---':
        return {'type': 'separator', 'content': ''}
    
    # Regular paragraph
    return {'type': 'paragraph', 'content': line}


def process_markdown_content(content: str) -> List[dict]:
    """Process markdown content and return structured elements."""
    lines = content.split('\n')
    elements = []
    in_code_block = False
    code_block_lines = []
    
    for line in lines:
        parsed = parse_markdown_line(line)
        
        if parsed['type'] == 'code_block_start':
            in_code_block = True
            continue
        elif parsed['type'] == 'code_block' and in_code_block:
            in_code_block = False
            if code_block_lines:
                elements.append({
                    'type': 'code',
                    'content': '\n'.join(code_block_lines)
                })
                code_block_lines = []
            continue
        
        if in_code_block:
            code_block_lines.append(line)
        else:
            elements.append(parsed)
    
    # Handle any remaining code block
    if code_block_lines:
        elements.append({
            'type': 'code',
            'content': '\n'.join(code_block_lines)
        })
    
    return elements


def add_formatted_text(paragraph, text: str):
    """Add formatted text to paragraph, handling markdown inline formatting."""
    # Handle bold **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            # Code text
            code_text = part[1:-1]
            run = paragraph.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Regular text
            paragraph.add_run(part)


def convert_markdown_to_docx(md_file: Path, output_file: Path):
    """Convert a markdown file to a Word document."""
    print(f"Converting {md_file.name} to {output_file.name}...")
    
    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Process markdown
    elements = process_markdown_content(content)
    
    # Create document
    doc = Document()
    setup_document_styles(doc)
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Process elements
    for element in elements:
        if element['type'] == 'empty':
            doc.add_paragraph()
        
        elif element['type'] == 'h1':
            p = doc.add_heading(element['content'], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        elif element['type'] == 'h2':
            p = doc.add_heading(element['content'], level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        elif element['type'] == 'h3':
            p = doc.add_heading(element['content'], level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        elif element['type'] == 'h4':
            p = doc.add_heading(element['content'], level=4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        elif element['type'] == 'bullet':
            p = doc.add_paragraph(element['content'], style='List Bullet')
            # Process inline formatting
            p.clear()
            add_formatted_text(p, element['content'])
        
        elif element['type'] == 'numbered':
            p = doc.add_paragraph(element['content'], style='List Number')
            # Process inline formatting
            p.clear()
            add_formatted_text(p, element['content'])
        
        elif element['type'] == 'code':
            # Code block
            p = doc.add_paragraph()
            p.style = 'No Spacing'
            run = p.add_run(element['content'])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            # Add gray background
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        
        elif element['type'] == 'separator':
            # Add horizontal rule
            p = doc.add_paragraph()
            p.add_run('_' * 80)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif element['type'] == 'paragraph':
            if element['content'].strip():
                p = doc.add_paragraph()
                add_formatted_text(p, element['content'])
        
        elif element['type'] == 'bold':
            p = doc.add_paragraph()
            run = p.add_run(element['content'])
            run.bold = True
    
    # Save document
    doc.save(output_file)
    print(f"✓ Saved: {output_file}")


def main():
    """Convert all markdown reports to Word documents."""
    base_dir = Path(__file__).parent
    
    # Reports to convert
    reports = [
        'WEEK_9_DEPLOYMENT_REPORT.md',
        'MODELING_EVALUATION_VALIDATION_REPORT.md',
        'WEEK_9_COMPLETION_SUMMARY.md',
        'APPLICATION_SUMMARY.md',
    ]
    
    # Create output directory
    output_dir = base_dir / 'outputs' / 'word_documents'
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print("Converting markdown reports to Word documents...\n")
    
    # Convert each report
    for report_name in reports:
        md_file = base_dir / report_name
        if md_file.exists():
            output_name = report_name.replace('.md', '.docx')
            output_file = output_dir / output_name
            try:
                convert_markdown_to_docx(md_file, output_file)
            except Exception as e:
                print(f"✗ Error converting {report_name}: {e}")
                import traceback
                traceback.print_exc()
        else:
            print(f"⚠ Warning: {report_name} not found")
    
    print(f"\n✓ Conversion complete! Documents saved to: {output_dir}")
    print(f"\nGenerated files:")
    for docx_file in sorted(output_dir.glob('*.docx')):
        size_kb = docx_file.stat().st_size / 1024
        print(f"  - {docx_file.name} ({size_kb:.1f} KB)")


if __name__ == '__main__':
    main()



